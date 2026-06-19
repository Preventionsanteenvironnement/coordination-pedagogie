#!/usr/bin/env python3
"""
Génère le sondage CPS en version texte accessible (machine braille).
Pas de tableaux, pas d'images, structure linéaire, navigation simple.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# === Styles de base accessibles ===
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(14)  # gros corps pour lisibilité

# Marges généreuses
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(17)

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)

def para(text=""):
    doc.add_paragraph(text)

def question(num, label, instr=None):
    p = doc.add_paragraph()
    r = p.add_run(f"Question {num} : {label}")
    r.bold = True
    if instr:
        para(instr)
    para("Ma réponse : ____________________________________________")
    para("")

# ==================================================================
# EN-TÊTE
# ==================================================================
h1("SONDAGE — BILAN COMPÉTENCES PSYCHOSOCIALES")
para("Terminale Bac Pro — Parcours éducatif de santé")
para("Version accessible (texte brut)")
para("")

h3("Identification")
para("Classe (TBP1, TBP2, TBP3 ou autre) : ____________")
para("Date de passation : ____________")
para("")

h3("Instructions")
para("Salut. Ce questionnaire sert à préparer les ateliers sur les compétences psychosociales : mieux se connaître, gérer ses émotions, communiquer, coopérer. Tes réponses sont anonymes — personne ne saura ce que tu as écrit personnellement.")
para("")
para("Pour chaque question, écris ta réponse sur la ligne prévue. Tu peux laisser certaines lignes vides si tu préfères ne pas répondre.")
para("")
para("Le sondage est divisé en 15 sections. Tu n'es pas obligé(e) de répondre à tout d'un seul coup.")
para("")

# ==================================================================
# SECTION 1 — Auto-positionnement (12 items, échelle 1-5)
# ==================================================================
h2("SECTION 1 sur 15 — Aujourd'hui, je me sens comment")
para("Pour chaque phrase, donne une note de 1 à 5 :")
para("1 = pas du tout à l'aise")
para("2 = un peu")
para("3 = moyen")
para("4 = bien")
para("5 = très à l'aise")
para("")

auto_questions = [
    ("1.1", "Parler de moi, raconter ce que je vis."),
    ("1.2", "Reconnaître mes qualités et mes points forts."),
    ("1.3", "Dire ce que je ressens (mes émotions)."),
    ("1.4", "Écouter vraiment quelqu'un sans l'interrompre."),
    ("1.5", "M'exprimer clairement à l'oral."),
    ("1.6", "Dire non sans agresser, savoir refuser."),
    ("1.7", "Travailler en groupe, coopérer."),
    ("1.8", "Demander de l'aide quand j'en ai besoin."),
    ("1.9", "Gérer mon stress, mes émotions difficiles."),
    ("1.10", "Prendre une décision importante."),
    ("1.11", "Me concentrer longtemps sur une tâche."),
    ("1.12", "Adapter mon langage selon la situation (pote, prof, pro)."),
]
for n, q in auto_questions:
    question(n, q, "Note de 1 à 5 :")

# ==================================================================
# SECTION 2 — Intérêt pour les thèmes (19 items, échelle 1-5)
# ==================================================================
h2("SECTION 2 sur 15 — Les thèmes qui m'intéressent")
para("Pour chaque thème, donne une note d'intérêt de 1 à 5 :")
para("1 = ça ne m'intéresse pas du tout")
para("5 = j'adorerais travailler là-dessus")
para("")

themes = [
    ("2.1", "Mieux me connaître (qui je suis, comment je fonctionne)"),
    ("2.2", "Comprendre mes qualités et mes points forts"),
    ("2.3", "Prendre confiance en moi"),
    ("2.4", "Comprendre mes besoins et mes valeurs"),
    ("2.5", "Réfléchir à mes objectifs personnels"),
    ("2.6", "Faire de meilleurs choix / décisions"),
    ("2.7", "Penser par moi-même, esprit critique"),
    ("2.8", "Mieux me concentrer, gérer mon attention"),
    ("2.9", "Comprendre mes émotions"),
    ("2.10", "Savoir nommer ce que je ressens"),
    ("2.11", "Gérer les moments où je suis mal à l'aise / stressé(e)"),
    ("2.12", "Mieux communiquer avec les autres"),
    ("2.13", "Apprendre à faire une demande claire"),
    ("2.14", "Apprendre à écouter vraiment quelqu'un"),
    ("2.15", "Comprendre les malentendus, désamorcer"),
    ("2.16", "Gérer les tensions et conflits dans un groupe"),
    ("2.17", "Mieux travailler en équipe"),
    ("2.18", "Aider, soutenir quelqu'un (comportements prosociaux)"),
    ("2.19", "Préparer la vie pro et les relations au travail"),
]
for n, q in themes:
    question(n, q, "Note de 1 à 5 :")

# ==================================================================
# SECTION 3 — Situations où progresser (5 max parmi 19)
# ==================================================================
h2("SECTION 3 sur 15 — Situations où je veux être plus à l'aise")
para("Cite au maximum 5 situations parmi les suivantes. Écris uniquement les numéros (ex : 3, 7, 12).")
para("")
situations = [
    ("1", "Parler devant la classe"),
    ("2", "Passer un oral (CCF, examen)"),
    ("3", "Gérer le stress avant une épreuve"),
    ("4", "Garder confiance pendant une difficulté"),
    ("5", "Demander de l'aide"),
    ("6", "Dire non correctement"),
    ("7", "Expliquer ce que je ressens"),
    ("8", "Éviter de m'énerver trop vite"),
    ("9", "Mieux comprendre les autres"),
    ("10", "Mieux travailler en groupe"),
    ("11", "Gérer une critique"),
    ("12", "Gérer un conflit"),
    ("13", "Me motiver quand je n'ai pas envie"),
    ("14", "Préparer mon avenir"),
    ("15", "Réussir en stage"),
    ("16", "Communiquer avec un tuteur ou un employeur"),
    ("17", "Prendre une décision importante"),
    ("18", "Mieux organiser mes idées"),
    ("19", "Me concentrer plus longtemps"),
]
for n, q in situations:
    para(f"{n}. {q}")
para("")
para("Mes 5 choix (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 4 — Difficultés vécues (à cocher / lister)
# ==================================================================
h2("SECTION 4 sur 15 — Ce qui me pose vraiment problème")
para("Liste les numéros de toutes les situations qui te ressemblent, même un peu (ex : 1, 5, 8, 11).")
para("")
difficultes = [
    ("1", "Quand je dois prendre la parole, je bloque, je rougis, je tremble."),
    ("2", "Quand quelqu'un m'énerve, je m'emporte et je le regrette après."),
    ("3", "Face à un adulte (prof, employeur), je me tais même si je ne suis pas d'accord."),
    ("4", "Quand mes potes m'entraînent, je n'ose pas refuser."),
    ("5", "En entretien (PFMP, embauche), je perds mes mots."),
    ("6", "Quand on me critique, je me ferme ou je contre-attaque."),
    ("7", "Je me sens jugé(e) dès que j'ouvre la bouche."),
    ("8", "Je n'arrive pas à dire ce que je ressens."),
    ("9", "Je coupe la parole sans m'en rendre compte."),
    ("10", "Je dis ce que les autres veulent entendre, pas ce que je pense."),
    ("11", "Mon langage de tous les jours me pose problème en stage / au travail."),
    ("12", "J'ai du mal à comprendre les sous-entendus."),
    ("13", "Je n'arrive pas à regarder dans les yeux quand je parle."),
    ("14", "Au téléphone, je suis encore plus mal à l'aise."),
    ("15", "Mes messages écrits (SMS, mails) sont mal pris."),
    ("16", "Quand je suis stressé(e), je ne sais plus parler."),
]
for n, q in difficultes:
    para(f"{n}. {q}")
para("")
para("Mes choix (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 5 — Vote ateliers (5 dans l'ordre)
# ==================================================================
h2("SECTION 5 sur 15 — Mes 5 ateliers prioritaires")
para("Voici 15 propositions d'ateliers. Choisis tes 5 préférés et classe-les dans l'ordre (du plus important au moins important). Écris les numéros séparés par des virgules (ex : 5, 1, 12, 3, 9).")
para("")
ateliers = [
    ("1", "Prendre la parole devant un groupe (gérer le trac, la voix, la posture)"),
    ("2", "Écouter pour vraiment comprendre (écoute active, reformulation)"),
    ("3", "Dire non sans casser la relation (assertivité, oser refuser)"),
    ("4", "Gérer un conflit sans s'énerver (désamorcer, négocier)"),
    ("5", "Communiquer en milieu pro (entretien, stage, employeur)"),
    ("6", "Adapter mon langage selon les situations (codes sociaux)"),
    ("7", "Faire passer un message clair (message-je, demande directe)"),
    ("8", "Convaincre sans manipuler (argumenter, négocier honnêtement)"),
    ("9", "Gérer mes émotions pendant que je parle (stress, colère, peur)"),
    ("10", "Décoder le non-verbal (gestes, regard, ton, silences)"),
    ("11", "Sortir d'une dispute en famille ou entre amis"),
    ("12", "Communication en ligne (messages, réseaux, mails pro)"),
    ("13", "Renforcer ma confiance en moi"),
    ("14", "Mieux travailler en équipe (coopération, répartition des rôles)"),
    ("15", "Mieux me connaître (forces, valeurs, fonctionnement)"),
]
for n, q in ateliers:
    para(f"{n}. {q}")
para("")
para("Mes 5 choix dans l'ordre (du plus important au moins important) : ____________________________________________")
para("")

# ==================================================================
# SECTION 6 — Activités préférées (5 max)
# ==================================================================
h2("SECTION 6 sur 15 — Activités qui me donnent envie")
para("Choisis maximum 5 types d'activités. Écris les numéros.")
para("")
activites = [
    ("1", "Jeux de rôle"),
    ("2", "Débats"),
    ("3", "Quiz interactifs"),
    ("4", "Jeux en équipe"),
    ("5", "Mises en situation"),
    ("6", "Défis courts (5-10 min)"),
    ("7", "Vidéos puis discussion"),
    ("8", "Cartes à manipuler"),
    ("9", "Situations de stage ou de travail"),
    ("10", "Situations de la vie réelle"),
    ("11", "Temps de réflexion personnelle"),
    ("12", "Activités créatives (dessin, écriture)"),
    ("13", "Échanges en petit groupe"),
    ("14", "Analyse de scènes ou dialogues"),
    ("15", "Mini-projets sur plusieurs séances"),
    ("16", "Exercices de respiration / concentration"),
    ("17", "Cercle de parole avec règles claires"),
    ("18", "S'enregistrer / se filmer"),
]
for n, q in activites:
    para(f"{n}. {q}")
para("")
para("Mes 5 choix (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 7 — Modalités de travail
# ==================================================================
h2("SECTION 7 sur 15 — Comment je préfère travailler")
para("")
h3("Question 7.1 : Configuration préférée (plusieurs choix possibles)")
para("1. Seul(e)")
para("2. En binôme")
para("3. En petit groupe (3-4)")
para("4. En classe entière")
para("5. Ça dépend du thème")
para("Mes choix (numéros) : ____________________________________________")
para("")
h3("Question 7.2 : Quand on travaille en groupe, je préfère (un seul choix)")
para("1. Choisir mon groupe")
para("2. Que le prof fasse les groupes")
para("3. Tirage au hasard")
para("4. Varier les groupes à chaque fois")
para("Mon choix (un seul numéro) : ____________")
para("")

# ==================================================================
# SECTION 8 — Guidage
# ==================================================================
h2("SECTION 8 sur 15 — Le niveau de guidage dont j'ai besoin")
para("")
h3("Question 8.1 : Pour être à l'aise, j'aime (plusieurs choix possibles)")
para("1. Consignes détaillées")
para("2. Consignes simples et courtes")
para("3. Un exemple avant de commencer")
para("4. Une fiche-guide étape par étape")
para("5. Une activité libre avec peu de consignes")
para("6. Un modèle à compléter")
para("7. Une correction collective à la fin")
para("Mes choix (numéros) : ____________________________________________")
para("")
h3("Question 8.2 : Un petit livret personnel pour suivre mes progrès m'aiderait (un seul choix)")
para("1. Oui, beaucoup")
para("2. Oui, un peu")
para("3. Non, pas vraiment")
para("4. Seulement pour certaines activités")
para("Mon choix : ____________")
para("")

# ==================================================================
# SECTION 9 — Format de séance
# ==================================================================
h2("SECTION 9 sur 15 — Format de séance qui me convient")
para("")
h3("Question 9.1 : Je préfère (un seul choix)")
para("1. Ateliers d'1 heure")
para("2. Ateliers de 2 heures")
para("3. Ateliers courts mais fréquents")
para("4. Ateliers plus longs mais moins fréquents")
para("5. Journée intensive (immersion)")
para("Mon choix : ____________")
para("")
h3("Question 9.2 : Dans une séance, je préfère (un seul choix)")
para("1. Une seule grande activité")
para("2. Plusieurs petites activités")
para("3. Une partie explication puis une activité")
para("4. Une activité puis une discussion")
para("5. Un défi puis un bilan")
para("Mon choix : ____________")
para("")

# ==================================================================
# SECTION 10 — Cadre et sécurité psychologique (5 max)
# ==================================================================
h2("SECTION 10 sur 15 — Le cadre dont j'ai besoin pour participer sereinement")
para("Choisis au maximum 5 conditions. Écris les numéros.")
para("")
cadre = [
    ("1", "Que personne ne se moque"),
    ("2", "Que les réponses restent respectées (pas de jugement)"),
    ("3", "Que je ne sois pas obligé(e) de parler"),
    ("4", "Que les consignes soient claires"),
    ("5", "Que le prof cadre bien le groupe"),
    ("6", "Que l'ambiance soit calme"),
    ("7", "Que l'ambiance soit dynamique"),
    ("8", "Que les activités soient utiles"),
    ("9", "Que les sujets restent concrets"),
    ("10", "Que chacun puisse donner son avis"),
    ("11", "Que les exemples parlent de situations réelles"),
    ("12", "Que les activités ne soient pas trop personnelles"),
]
for n, q in cadre:
    para(f"{n}. {q}")
para("")
para("Mes 5 choix maximum (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 11 — Rapport à l'oral
# ==================================================================
h2("SECTION 11 sur 15 — Mon rapport à la prise de parole")
para("")
h3("Question 11.1 : Pendant les ateliers, parler devant les autres, pour moi c'est (un seul choix)")
para("1. Facile")
para("2. Possible si je suis à l'aise")
para("3. Difficile")
para("4. Très difficile")
para("5. Je préfère parler en petit groupe")
para("Mon choix : ____________")
para("")
h3("Question 11.2 : Je préfère participer (plusieurs choix possibles)")
para("1. À l'oral")
para("2. Par écrit")
para("3. Par vote anonyme")
para("4. Avec le téléphone (quiz, sondage)")
para("5. En groupe")
para("6. Ça dépend du sujet")
para("Mes choix (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 12 — Sujets sensibles
# ==================================================================
h2("SECTION 12 sur 15 — Sujets à aborder avec prudence")
para("Attention : tu ne seras jamais obligé(e) de raconter quelque chose de personnel. Cette section sert juste à savoir comment cadrer les ateliers.")
para("")
para("Coche les sujets pour lesquels tu préfères qu'on prenne des précautions (ou 12 si rien ne te dérange).")
para("")
sujets = [
    ("1", "Stress, anxiété"),
    ("2", "Confiance en soi"),
    ("3", "Émotions intimes"),
    ("4", "Famille"),
    ("5", "Relations amicales"),
    ("6", "Relations amoureuses"),
    ("7", "Conflits"),
    ("8", "Harcèlement"),
    ("9", "Avenir / orientation"),
    ("10", "Stage / PFMP"),
    ("11", "Oral / prise de parole"),
    ("12", "Aucun sujet ne me pose problème"),
]
for n, q in sujets:
    para(f"{n}. {q}")
para("")
para("Mes choix (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 13 — Situations pro (5 max)
# ==================================================================
h2("SECTION 13 sur 15 — Ateliers utiles pour ma vie pro")
para("Choisis au maximum 5 ateliers. Écris les numéros.")
para("")
pro = [
    ("1", "Se présenter à un employeur"),
    ("2", "Communiquer avec un tuteur"),
    ("3", "Gérer une remarque"),
    ("4", "Demander une explication"),
    ("5", "Dire qu'on n'a pas compris"),
    ("6", "Travailler avec des collègues"),
    ("7", "Gérer un désaccord"),
    ("8", "Être ponctuel et fiable"),
    ("9", "Garder son calme"),
    ("10", "Prendre une décision"),
    ("11", "Valoriser ses qualités"),
    ("12", "Préparer un entretien"),
    ("13", "Préparer l'oral du bac"),
    ("14", "Expliquer son projet"),
    ("15", "Comprendre les codes du monde pro"),
]
for n, q in pro:
    para(f"{n}. {q}")
para("")
para("Mes 5 choix maximum (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 14 — Motivations et freins
# ==================================================================
h2("SECTION 14 sur 15 — Motivations et freins")
para("")
h3("Question 14.1 : Ce qui me donne envie de participer (plusieurs choix possibles)")
motivations = [
    ("1", "C'est concret"),
    ("2", "C'est utile pour ma vie"),
    ("3", "C'est utile pour le bac"),
    ("4", "C'est utile pour le travail"),
    ("5", "Il y a du jeu"),
    ("6", "Il y a du mouvement"),
    ("7", "Je peux donner mon avis"),
    ("8", "Je peux choisir"),
    ("9", "Je comprends pourquoi on le fait"),
    ("10", "L'ambiance est bonne"),
    ("11", "Ce n'est pas trop long"),
    ("12", "Je réussis quelque chose"),
]
for n, q in motivations:
    para(f"{n}. {q}")
para("Mes choix (numéros) : ____________________________________________")
para("")
h3("Question 14.2 : Ce qui peut me bloquer (plusieurs choix possibles)")
freins = [
    ("1", "Peur du regard des autres"),
    ("2", "Peur de parler"),
    ("3", "Peur de me tromper"),
    ("4", "Manque d'intérêt"),
    ("5", "Consignes trop longues"),
    ("6", "Trop d'écrit"),
    ("7", "Trop d'oral"),
    ("8", "Groupe trop bruyant"),
    ("9", "Activité trop personnelle"),
    ("10", "Difficulté à me concentrer"),
    ("11", "Impression que ça ne sert à rien"),
]
for n, q in freins:
    para(f"{n}. {q}")
para("Mes choix (numéros) : ____________________________________________")
para("")

# ==================================================================
# SECTION 15 — Verbatim et profil
# ==================================================================
h2("SECTION 15 sur 15 — Pour finir, à toi la parole")
para("Les 4 premières questions sont libres : tu peux laisser vide si tu veux.")
para("")

h3("Question 15.1")
para("Raconte une situation où la communication a mal tourné et que tu aimerais savoir gérer la prochaine fois.")
para("Ma réponse : ____________________________________________")
para("____________________________________________")
para("____________________________________________")
para("")
h3("Question 15.2")
para("Un truc précis que tu veux absolument apprendre dans ces ateliers :")
para("Ma réponse : ____________________________________________")
para("____________________________________________")
para("")
h3("Question 15.3")
para("Si tu pouvais proposer un atelier, ce serait sur quoi ?")
para("Ma réponse : ____________________________________________")
para("____________________________________________")
para("")
h3("Question 15.4")
para("Qu'est-ce qui ferait qu'un atelier soit vraiment réussi pour toi ?")
para("Ma réponse : ____________________________________________")
para("____________________________________________")
para("")

h3("Question 15.5 : En communication, je me considère plutôt (un seul choix)")
para("1. Réservé(e)")
para("2. Dans la moyenne")
para("3. Expansif(ve)")
para("4. Ça dépend des situations")
para("Mon choix : ____________")
para("")

h3("Question 15.6 : Après le bac, mon projet principal (un seul choix)")
para("1. Entrer dans la vie active")
para("2. Poursuivre des études")
para("3. Alternance / apprentissage")
para("4. Je ne sais pas encore")
para("Mon choix : ____________")
para("")

h3("Question 15.7 : Ma classe (un seul choix)")
para("1. Terminale Bac Pro 1 (TBP1)")
para("2. Terminale Bac Pro 2 (TBP2)")
para("3. Terminale Bac Pro 3 (TBP3)")
para("4. Autre")
para("Mon choix : ____________")
para("")

doc.add_page_break()

# ==================================================================
# PROMPT IA POUR LE PROF (page séparée)
# ==================================================================
h1("INSTRUCTIONS POUR LE PROFESSEUR")
para("Ne pas donner cette page à l'élève. Elle sert à convertir ses réponses en JSON via une intelligence artificielle.")
para("")

h2("Étape 1 : Récupérer les réponses de l'élève")
para("L'élève remplit le questionnaire ci-dessus sur sa machine braille. Récupère son texte (par mail, fichier, dictée…).")
para("")

h2("Étape 2 : Prompt à coller dans Claude / ChatGPT")
para("Copie-colle le bloc ci-dessous, puis ajoute les réponses de l'élève à la fin, puis envoie.")
para("")

prompt = """
Tu vas recevoir les réponses d'un élève à un sondage CPS de 15 sections. Convertis-les en un JSON STRICTEMENT au format suivant (ne change aucun nom de champ).

Format attendu :
{
  "auto_eval": { "ae1": <1-5>, "ae2": <1-5>, ..., "ae12": <1-5> },
  "themes": { "th1": <1-5>, "th2": <1-5>, ..., "th19": <1-5> },
  "situations_progresser": [ "s3", "s7", ... ],
  "difficultes": [ "d1", "d5", ... ],
  "vote_ateliers": [ "a5", "a1", "a12", "a3", "a9" ],
  "activites_preferees": [ "ac1", "ac5", ... ],
  "modalites_travail": [ "mo2", "mo3", ... ],
  "groupes_pref": "g1",
  "guidage_pref": [ "gu1", "gu3", ... ],
  "livret_pref": "l1",
  "duree_pref": "du1",
  "structure_pref": "st3",
  "cadre_besoins": [ "ca1", "ca4", ... ],
  "oral_aisance": "o3",
  "oral_modes": [ "op2", "op4", ... ],
  "sujets_sensibles": [ "se4", "se7" ],
  "situations_pro": [ "pr1", "pr5", ... ],
  "motivations": [ "mt1", "mt9", ... ],
  "freins": [ "fr1", "fr3", ... ],
  "verbatim_situation_difficile": "texte libre élève section 15.1",
  "verbatim_objectif": "texte libre élève section 15.2",
  "verbatim_atelier_propose": "texte libre élève section 15.3",
  "verbatim_atelier_reussi": "texte libre élève section 15.4",
  "profil": "reserve | moyenne | expansif | variable",
  "projet_post_bac": "vie_active | etudes | alternance | indecis",
  "classe": "TBP1 | TBP2 | TBP3 | autre",
  "source": "version_accessible_braille"
}

Règles de conversion :
- Section 1 et 2 : les numéros (1-5) deviennent des entiers, préfixe "ae" pour section 1 (ae1 à ae12), "th" pour section 2 (th1 à th19).
- Sections multi-choix : retourne UNIQUEMENT un tableau avec les codes (préfixe : s/d/a/ac/mo/gu/ca/op/se/pr/mt/fr) suivi du numéro de l'item. Exemple : si l'élève a choisi situations 3, 7, 12 → ["s3","s7","s12"].
- Section 5 (vote ateliers) : l'ordre compte. Premier choix = priorité 1. Préfixe "a". Exemple "5, 1, 12, 3, 9" → ["a5","a1","a12","a3","a9"].
- Choix unique (radio) : code seul, sans tableau. Exemples : groupes_pref "1" → "g1", livret_pref "2" → "l2", duree_pref "3" → "du3", structure_pref "4" → "st4", oral_aisance "3" → "o3", profil section 15.5 "1" → "reserve" (et non "p1"), projet_post_bac "2" → "etudes", classe "1" → "TBP1".
- Si une réponse est vide ou non interprétable, mets [] pour les tableaux, null pour les choix uniques, "" pour les verbatim.
- Ne mets PAS de champ "auto_eval" partiel : si l'élève n'a pas tout rempli, mets quand même les ae1..ae12 avec null pour les manquants.

Réponds UNIQUEMENT par le JSON, sans aucun commentaire ni balise markdown. Le JSON doit être directement parsable.

Voici les réponses de l'élève :
[COLLE ICI LES RÉPONSES DE L'ÉLÈVE]
"""
para(prompt)
para("")

h2("Étape 3 : Importer le JSON dans le dashboard")
para("Une fois le JSON généré par l'IA, copie-le.")
para("Va dans le Cockpit PSE, ouvre 'Sondage CPS', clique sur le bouton '📥 Importer une réponse (JSON)' en haut.")
para("Colle le JSON et valide. La réponse est ajoutée aux statistiques de la classe.")
para("")

import os
out = os.path.expanduser("~/Documents/Bilan_competences_accessible.docx")
doc.save(out)
print(f"✅ Word généré (local, hors GitHub) : {out}")
